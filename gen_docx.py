#!/usr/bin/env python3
"""Generate cv-harvard.docx (EN) and cv-harvard-es.docx (ES) from content."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

def add_bottom_border(paragraph, sz=12, color="1A1A1A"):
    """Add a bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): str(sz),
        qn('w:space'): '1',
        qn('w:color'): color,
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_run_font(run, name='Times New Roman', size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): name,
        qn('w:hAnsi'): name,
        qn('w:cs'): name,
    })
    rPr.append(rFonts)

def section_heading(doc, text_es, text_en, lang):
    text = text_es if lang == 'es' else text_en
    p = doc.add_paragraph()
    add_bottom_border(p)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, bold_parts=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    # Remove default bullet
    p.style.font.size = Pt(10)
    # Clear and add manually
    p.clear()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_run_font(run, size=10)
    return p

def entry_header(doc, org, role, location, date, lang):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(org)
    set_run_font(run, size=10.5, bold=True)
    
    run2 = p.add_run(f" — {role}")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    
    # Date on right — use tab stop
    tab_stops = p.paragraph_format.tab_stops
    # We'll just append with spacing
    run3 = p.add_run(f"\t{date}")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    return p

def skills_row(doc, label, value, lang):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    
    run = p.add_run(label)
    set_run_font(run, size=10, bold=True)
    
    run2 = p.add_run(value)
    set_run_font(run2, size=10)
    return p

def cert_item(doc, name, issuer, date, lang):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    
    run = p.add_run(name)
    set_run_font(run, size=9.5, bold=True)
    
    run2 = p.add_run(f" — {issuer}")
    set_run_font(run2, size=9.5)
    
    if date:
        run3 = p.add_run(f", {date}")
        set_run_font(run3, size=9.5)
    
    return p

# ============================================================
# ENGLISH VERSION
# ============================================================
def generate_en():
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # HEADER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("ANTONIO CORTAZAR JIMENEZ")
    set_run_font(run, size=18, bold=True)
    add_bottom_border(p, sz=16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run("Cancún, Mexico | 81 3511 2848 | hola@antoniocortazar.dev")
    set_run_font(run, size=9.5, color=(51,51,51))
    run2 = p2.add_run("\nlinkedin.com/in/antoniocortazar | antoniocortazar.dev | github.com/Antony-potato")
    set_run_font(run2, size=9.5, color=(51,51,51))

    # PROFILE
    section_heading(doc, "Professional Summary", "", "en")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(
        "Full Stack Developer and Software Development Engineering student with C1 English proficiency. "
        "Specialized in building scalable web solutions and immersive interfaces using modern technologies "
        "like React, Next.js, and Astro, integrated with cloud infrastructures (AWS, Docker). "
        "Results-oriented and committed to delivering efficient code. I combine a strong technical foundation "
        "with excellent interpersonal skills to collaborate in multidisciplinary teams, digitize corporate "
        "processes, and create technological tools that deliver real business value and enhance user experience."
    )
    set_run_font(run, size=10)

    # EDUCATION
    section_heading(doc, "Education", "", "en")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Universidad Tecmilenio")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Cancún, Mexico")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\t2023 – Expected: May 2027")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run("B.S. in Software Development Engineering")
    set_run_font(run, size=10, italic=True)
    
    bullet(doc, "Relevant Coursework: Data Structures, Object-Oriented Programming (Java), Full Stack Web Development, Relational Databases, Infrastructure as Code.")

    # EXPERIENCE
    section_heading(doc, "Experience", "", "en")
    
    # Onix Living
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Onix Living")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — IT Assistant & Web Developer")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\tAug 2025 – Present")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Containerized and deployed internal web applications using Docker Compose on AWS Lightsail, optimizing release times and ensuring consistency across environments.")
    bullet(doc, "Built an internal IT operations dashboard with Next.js and Gridstack, providing drag-and-drop widgets for real-time monitoring of asset status and support tickets.")
    bullet(doc, "Implemented Snipe-IT asset management system to track 100+ fixed assets; automated provisioning and decommissioning workflows with web forms and operational dashboards.")
    bullet(doc, "Deployed static web applications on AWS and managed multimedia asset storage via Amazon S3 for company projects.")
    bullet(doc, "Provided Tier 1 and Tier 2 technical support for hardware and software across office infrastructure; trained 25+ employees on AI tools and Zoho One automation workflows.")
    
    # Office Depot
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Office Depot")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Technology Department Supervisor")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\tJun 2024 – Aug 2025")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Led a team of 5 associates in the technology department, providing training and daily operational guidance.")
    bullet(doc, "Drove technology product warranty sales and exceeded quarterly KPI targets through customer-facing consultations on hardware, software, and electronics.")
    bullet(doc, "Managed department KPIs including warranty sales targets, customer satisfaction scores, and inventory accuracy.")
    bullet(doc, "Provided technical troubleshooting and product demonstrations to customers.")

    # PROJECTS
    section_heading(doc, "Projects", "", "en")
    
    # Koa Towers
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Koa Towers & Aldea Savia — Interactive Real Estate Platforms")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Onix Living")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\t2025")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Architected and built static landing pages with Astro 6 (SSG) to digitize the presentation of two real estate projects (previously distributed in disorganized PDF files), dramatically improving SEO positioning and achieving sub-second page load times.")
    bullet(doc, "Built immersive interfaces with interactive 3D maps and scroll-triggered animations, synchronizing dynamic data via Google Sheets API to visualize inventory of 180+ units in real time.")
    bullet(doc, "Configured automated CI/CD pipelines with GitHub Actions for scheduled rebuilds every 12 hours, and deployment to custom domain.")
    
    # Pocky
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Pocky — Cross-Platform Virtual Pet PWA")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Personal")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Developed a shared virtual pet application with real-time synchronization via Firebase across iOS (Safari PWA) and Android (Capacitor) without App Store or Google Play distribution.")
    bullet(doc, "Implemented 24/7 pet degradation system, animated characters, and Web Push notifications for cross-platform engagement.")
    
    # Microservices
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Microservices Architecture — Docker + Kubernetes + Helm")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Academic")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Designed a full microservices stack with Node.js backend, nginx reverse-proxy frontend, and MySQL database.")
    bullet(doc, "Configured Kubernetes deployments via Helm charts (Secrets, ConfigMaps, PVCs, Services) and automated CI/CD pipeline with GitHub Actions for Docker Hub publishing.")

    # SKILLS
    section_heading(doc, "Skills", "", "en")
    
    skills_row(doc, "Frontend:", "Next.js, React, Astro, Vue 3, Tailwind CSS, Framer Motion, Swiper", "en")
    skills_row(doc, "Languages:", "JavaScript, TypeScript, HTML5, CSS3, Java, PHP, SQL", "en")
    skills_row(doc, "Backend:", "Node.js, Spring Boot, Laravel", "en")
    skills_row(doc, "Infrastructure & DB:", "AWS, Amazon S3, Docker, Kubernetes, Helm, GitHub Actions, MySQL", "en")
    skills_row(doc, "Other:", "Firebase (Realtime, Admin), PWA, Capacitor, Git, Snipe-IT, Figma", "en")

    # CERTIFICATIONS
    section_heading(doc, "Certifications", "", "en")
    
    cert_item(doc, "EF SET English — C1 Advanced (67/100)", "EF SET", "Jul 2025", "en")
    cert_item(doc, "AWS Academy Graduate — Data Engineering", "AWS", "Dec 2025", "en")
    cert_item(doc, "Figma Avanzado", "LinkedIn Learning", "", "en")
    cert_item(doc, "Networking Academy Learn-A-Thon 2024", "Cisco", "Jun 2024", "en")
    cert_item(doc, "Big Data & Hadoop Foundations (L1, L2, Spark)", "IBM", "Apr 2026", "en")

    doc.save("cv-harvard.docx")
    print("Saved cv-harvard.docx")

# ============================================================
# SPANISH VERSION
# ============================================================
def generate_es():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # HEADER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("ANTONIO CORTAZAR JIMENEZ")
    set_run_font(run, size=18, bold=True)
    add_bottom_border(p, sz=16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run("Cancún, México | 81 3511 2848 | hola@antoniocortazar.dev")
    set_run_font(run, size=9.5, color=(51,51,51))
    run2 = p2.add_run("\nlinkedin.com/in/antoniocortazar | antoniocortazar.dev | github.com/Antony-potato")
    set_run_font(run2, size=9.5, color=(51,51,51))

    # PERFIL
    section_heading(doc, "Perfil Profesional", "", "es")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(
        "Desarrollador Full Stack y estudiante de Ingeniería en Desarrollo de Software con nivel de inglés C1. "
        "Especializado en construir soluciones web escalables e interfaces inmersivas utilizando tecnologías modernas "
        "como React, Next.js y Astro, integradas con infraestructuras en la nube (AWS, Docker). "
        "Orientado a resultados y comprometido con la entrega de código eficiente. Combino una sólida base técnica "
        "con excelentes habilidades interpersonales para colaborar en equipos multidisciplinarios, digitalizar procesos "
        "corporativos y crear herramientas tecnológicas que aporten valor real al negocio y a la experiencia del usuario."
    )
    set_run_font(run, size=10)

    # EDUCACIÓN
    section_heading(doc, "Educación", "", "es")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Universidad Tecmilenio")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Cancún, México")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\t2023 – Esperado: Mayo 2027")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run("Ingeniería en Desarrollo de Software")
    set_run_font(run, size=10, italic=True)
    
    bullet(doc, "Materias relevantes: Estructuras de Datos, Programación Orientada a Objetos (Java), Desarrollo Web Full Stack, Bases de Datos Relacionales, Infraestructura como Código.")

    # EXPERIENCIA
    section_heading(doc, "Experiencia", "", "es")
    
    # Onix Living
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Onix Living")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Auxiliar de TI y Desarrollo Web")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\tAgo 2025 – Presente")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Containericé y desplegué aplicaciones web internas mediante Docker Compose en AWS Lightsail, optimizando los tiempos de liberación y asegurando la consistencia entre entornos.")
    bullet(doc, "Construí un dashboard interno de operaciones de TI con Next.js y Gridstack, proporcionando widgets drag-and-drop para monitoreo en tiempo real del estado de activos y tickets de soporte.")
    bullet(doc, "Implementé el sistema de gestión de activos Snipe-IT para rastrear 100+ activos fijos; automaticé flujos de aprovisionamiento y baja con formularios web y dashboards operativos.")
    bullet(doc, "Desplegué aplicaciones web estáticas en AWS y gestioné almacenamiento de assets multimedia vía Amazon S3 para proyectos de la empresa.")
    bullet(doc, "Brindé soporte técnico Nivel 1 y 2 para hardware y software en infraestructura de oficina; capacité a 25+ empleados en herramientas de IA y flujos de automatización de Zoho One.")
    
    # Office Depot
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Office Depot")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Supervisor de Departamento de Tecnología")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\tJun 2024 – Ago 2025")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Lideré un equipo de 5 colaboradores en el departamento de tecnología, brindando capacitación y guía operativa diaria.")
    bullet(doc, "Impulsé ventas de garantías de productos tecnológicos y superé objetivos trimestrales de KPIs mediante consultas a clientes sobre hardware, software y electrónica.")
    bullet(doc, "Gestioné KPIs del departamento incluyendo metas de ventas de garantías, satisfacción del cliente y precisión de inventario.")
    bullet(doc, "Brindé resolución de problemas técnicos y demostraciones de productos a clientes.")

    # PROYECTOS
    section_heading(doc, "Proyectos", "", "es")
    
    # Koa Towers
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Koa Towers y Aldea Savia — Plataformas Inmobiliarias Interactivas")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Onix Living")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    run3 = p.add_run("\t2025")
    set_run_font(run3, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Arquitecté y construí landing pages estáticas con Astro 6 (SSG) para digitalizar la presentación de dos proyectos inmobiliarios (anteriormente distribuidos en archivos PDF desorganizados), mejorando drásticamente el posicionamiento SEO y logrando tiempos de carga de página por debajo de un segundo.")
    bullet(doc, "Construí interfaces inmersivas mediante mapas interactivos 3D y animaciones al hacer scroll, sincronizando datos dinámicos vía API de Google Sheets para visualizar el inventario de 180+ unidades en tiempo real.")
    bullet(doc, "Configuré pipelines de CI/CD automatizados con GitHub Actions para rebuilds programados cada 12 horas, y despliegue a dominio personalizado.")
    
    # Pocky
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Pocky — Mascota Virtual PWA Multiplataforma")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Personal")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Desarrollé una aplicación de mascota virtual compartida con sincronización en tiempo real vía Firebase a través de iOS (Safari PWA) y Android (Capacitor) sin distribución por App Store ni Google Play.")
    bullet(doc, "Implementé sistema de degradación 24/7 de la mascota, personajes animados y notificaciones Web Push para engagement multiplataforma.")
    
    # Microservicios
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Arquitectura de Microservicios — Docker + Kubernetes + Helm")
    set_run_font(run, size=10.5, bold=True)
    run2 = p.add_run(" — Académico")
    set_run_font(run2, size=9.5, italic=True, color=(85,85,85))
    
    bullet(doc, "Diseñé un stack completo de microservicios con backend Node.js, frontend con reverse-proxy nginx y base de datos MySQL.")
    bullet(doc, "Configuré despliegues en Kubernetes mediante Helm charts (Secrets, ConfigMaps, PVCs, Services) y automaticé pipeline CI/CD con GitHub Actions para publicación en Docker Hub.")

    # HABILIDADES
    section_heading(doc, "Habilidades", "", "es")
    
    skills_row(doc, "Frontend:", "Next.js, React, Astro, Vue 3, Tailwind CSS, Framer Motion, Swiper", "es")
    skills_row(doc, "Lenguajes:", "JavaScript, TypeScript, HTML5, CSS3, Java, PHP, SQL", "es")
    skills_row(doc, "Backend:", "Node.js, Spring Boot, Laravel", "es")
    skills_row(doc, "Infraestructura y BD:", "AWS, Amazon S3, Docker, Kubernetes, Helm, GitHub Actions, MySQL", "es")
    skills_row(doc, "Otros:", "Firebase (Realtime, Admin), PWA, Capacitor, Git, Snipe-IT, Figma", "es")

    # CERTIFICACIONES
    section_heading(doc, "Certificaciones", "", "es")
    
    cert_item(doc, "EF SET English — C1 Advanced (67/100)", "EF SET", "Jul 2025", "es")
    cert_item(doc, "AWS Academy Graduate — Data Engineering", "AWS", "Dic 2025", "es")
    cert_item(doc, "Figma Avanzado", "LinkedIn Learning", "", "es")
    cert_item(doc, "Networking Academy Learn-A-Thon 2024", "Cisco", "Jun 2024", "es")
    cert_item(doc, "Big Data & Hadoop Foundations (L1, L2, Spark)", "IBM", "Abr 2026", "es")

    doc.save("cv-harvard-es.docx")
    print("Saved cv-harvard-es.docx")

generate_en()
generate_es()
